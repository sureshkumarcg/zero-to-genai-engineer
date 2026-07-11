"""Generate reproducible sample documents for the ingestion & chunking notebooks.

Creates, in this folder:
  - sample_report.pdf     text + a real table (the "parsing quality-loss" demo)
  - sample_page.html      headings, paragraphs, a list, and a table
  - sample_contract.docx  numbered contract clauses
  - sample_scanned.png    a memo rendered as an image (for OCR)

Run:  python3 _generate_sample_data.py
"""
from pathlib import Path

HERE = Path(__file__).resolve().parent


# ── 1. PDF with a table ──────────────────────────────────────────────────────
def make_pdf():
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle)

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(HERE / "sample_report.pdf"), pagesize=LETTER)
    story = []

    story.append(Paragraph("Northwind Logistics — Q2 2024 Operations Report", styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "This report summarizes regional shipment performance for the second "
        "quarter of 2024. Overall shipment volume grew 8.4% quarter over quarter, "
        "driven primarily by the West and South regions. On-time delivery held "
        "above target in every region except the Northeast, where a warehouse "
        "migration temporarily reduced throughput.", styles["BodyText"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Regional performance is broken out in the table below.",
                           styles["BodyText"]))
    story.append(Spacer(1, 12))

    data = [
        ["Region", "Shipments", "On-Time %", "Revenue ($M)"],
        ["Northwest", "128,400", "97.2%", "14.6"],
        ["West", "204,910", "98.1%", "23.9"],
        ["South", "176,220", "96.8%", "19.4"],
        ["Northeast", "92,180", "89.3%", "11.2"],
        ["Total", "601,710", "95.9%", "69.1"],
    ]
    table = Table(data, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#faf6f0")),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
    ]))
    story.append(table)
    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "Action items: (1) complete the Northeast warehouse migration by July 15; "
        "(2) reallocate two delivery routes from Northwest to West to match demand; "
        "(3) review carrier contracts ahead of the Q3 peak season.", styles["BodyText"]))
    doc.build(story)
    print("wrote sample_report.pdf")


# ── 2. HTML page ─────────────────────────────────────────────────────────────
def make_html():
    html = """<!DOCTYPE html>
<html lang="en">
<head><meta charset="utf-8"><title>Northwind Logistics — Support FAQ</title></head>
<body>
  <h1>Northwind Logistics — Support FAQ</h1>
  <p>Answers to common questions about shipping, tracking, and returns.</p>

  <h2>Shipping</h2>
  <p>Standard shipping takes 3–5 business days. Expedited shipping takes 1–2
  business days and is available in the West and Northwest regions.</p>
  <ul>
    <li>Free standard shipping on orders over $75.</li>
    <li>Expedited shipping is a flat $12.</li>
    <li>We do not ship to PO boxes.</li>
  </ul>

  <h2>Tracking</h2>
  <p>Every shipment includes a tracking number sent by email once the order
  leaves the warehouse. Tracking updates every 6 hours.</p>

  <h2>Returns</h2>
  <p>Returns are accepted within 30 days of delivery. Refunds are processed to
  the original payment method within 5 business days of receipt.</p>
  <table border="1">
    <tr><th>Item type</th><th>Return window</th><th>Restocking fee</th></tr>
    <tr><td>Standard goods</td><td>30 days</td><td>None</td></tr>
    <tr><td>Perishable</td><td>Not returnable</td><td>&mdash;</td></tr>
    <tr><td>Bulk / freight</td><td>14 days</td><td>10%</td></tr>
  </table>
</body>
</html>
"""
    (HERE / "sample_page.html").write_text(html, encoding="utf-8")
    print("wrote sample_page.html")


# ── 3. Word contract ─────────────────────────────────────────────────────────
def make_docx():
    from docx import Document

    doc = Document()
    doc.add_heading("Master Services Agreement", level=0)
    doc.add_paragraph(
        "This Master Services Agreement (\"Agreement\") is entered into between "
        "Northwind Logistics, Inc. (\"Provider\") and the Client, effective as of "
        "the date of last signature.")

    clauses = [
        ("1. Scope of Services",
         "Provider shall deliver freight logistics and warehousing services as "
         "described in each executed Statement of Work."),
        ("2. Term",
         "This Agreement remains in effect for an initial term of twelve (12) "
         "months and renews automatically for successive twelve-month terms unless "
         "either party gives sixty (60) days written notice of non-renewal."),
        ("3. Fees and Payment",
         "Client shall pay all undisputed invoices within thirty (30) days of "
         "receipt. Late payments accrue interest at 1.5% per month."),
        ("4. Termination",
         "Either party may terminate this Agreement for material breach if the "
         "breach remains uncured thirty (30) days after written notice. Provider "
         "may suspend services immediately for non-payment exceeding sixty (60) days."),
        ("5. Remote-Work Reimbursement (HR-2024-17 §7.3)",
         "Employees working remotely at least three (3) days per week are eligible "
         "for a home-internet stipend of $45 per month, reimbursed via the monthly "
         "expense report. Contractors are not eligible."),
        ("6. Limitation of Liability",
         "Neither party's aggregate liability shall exceed the fees paid under this "
         "Agreement in the twelve (12) months preceding the claim."),
    ]
    for title, body in clauses:
        doc.add_heading(title, level=1)
        doc.add_paragraph(body)

    doc.save(str(HERE / "sample_contract.docx"))
    print("wrote sample_contract.docx")


# ── 4. Scanned-style image (for OCR) ─────────────────────────────────────────
def make_scanned_png():
    from PIL import Image, ImageDraw, ImageFont

    lines = [
        "INTERNAL MEMO",
        "",
        "TO:   All warehouse staff",
        "FROM: Operations, Northwind Logistics",
        "DATE: 2024-06-14",
        "RE:   Incident INC-5521",
        "",
        "The claims-search service had a 22-minute outage on",
        "2024-06-14 caused by a Qdrant snapshot restore that",
        "exhausted disk. Fix: volume size increased and a disk-",
        "usage alert added at 80%. Owner: platform team.",
    ]
    W, H = 900, 520
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Courier New.ttf", 24)
        title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 30)
    except OSError:
        font = ImageFont.load_default()
        title_font = font

    y = 30
    for i, line in enumerate(lines):
        draw.text((40, y), line, fill="black", font=title_font if i == 0 else font)
        y += 40 if i == 0 else 34
    # a light border to look like a scan
    draw.rectangle([5, 5, W - 6, H - 6], outline="#cccccc", width=2)
    img.save(str(HERE / "sample_scanned.png"))
    print("wrote sample_scanned.png")


if __name__ == "__main__":
    make_pdf()
    make_html()
    make_docx()
    make_scanned_png()
    print("Done. Sample data in", HERE)
